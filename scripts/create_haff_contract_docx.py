from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = "output/docx/HAFF_Leisure_Club_System_Agreement.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(table, top=80, start=120, bottom=80, end=120):
    tbl_pr = table._tbl.tblPr
    tbl_cell_mar = tbl_pr.find(qn("w:tblCellMar"))
    if tbl_cell_mar is None:
        tbl_cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(tbl_cell_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tbl_cell_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tbl_cell_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_grid = table._tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)

    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")


def set_borders(table, color="B8C2CC", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        border = borders.find(qn(tag))
        if border is None:
            border = OxmlElement(tag)
            borders.append(border)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), size)
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)


def style_run(run, bold=False, size=None, color=None):
    run.bold = bold
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_clause(doc, heading, paragraphs):
    doc.add_heading(heading, level=1)
    for para in paragraphs:
        if isinstance(para, list):
            for item in para:
                doc.add_paragraph(item, style="List Bullet")
        else:
            doc.add_paragraph(para)


def add_signature_block(doc):
    doc.add_heading("13. Signatures", level=1)
    doc.add_paragraph(
        "By signing below, the parties confirm that they have read, understood, and agreed to this Agreement."
    )
    table = doc.add_table(rows=5, cols=2)
    table.style = "Table Grid"
    set_table_width(table, [4560, 4560])
    set_cell_margins(table)
    set_borders(table)
    labels = [
        ("For the Client", "For the Developer / Service Provider"),
        ("Name: ______________________________", "Name: ______________________________"),
        ("Title: _______________________________", "Title: _______________________________"),
        ("Signature: ___________________________", "Signature: ___________________________"),
        ("Date: _______________________________", "Date: _______________________________"),
    ]
    for row_idx, row in enumerate(table.rows):
        for col_idx, text in enumerate(labels[row_idx]):
            cell = row.cells[col_idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(text)
            if row_idx == 0:
                style_run(r, bold=True, color="1F4D78")
                set_cell_shading(cell, "F2F4F7")


def build():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    title_style = styles["Title"]
    title_style.font.name = "Calibri"
    title_style.font.size = Pt(18)
    title_style.font.bold = True
    title_style.font.color.rgb = RGBColor.from_string("0B2545")
    title_style.paragraph_format.space_after = Pt(6)

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("HAFF Leisure Club System Agreement")
    style_run(run, size=9, color="667085")

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("Confidential - For discussion and signature")

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("SYSTEM DEVELOPMENT AND IMPLEMENTATION AGREEMENT")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("HAFF Leisure Club Court Management System")
    style_run(r, bold=True, size=12, color="1F4D78")

    meta = doc.add_table(rows=5, cols=2)
    meta.style = "Table Grid"
    set_table_width(meta, [2340, 6780])
    set_cell_margins(meta)
    set_borders(meta)
    rows = [
        ("Effective Date", "______________________________"),
        ("Client", "HAFF Leisure Club, Cadiz City, Philippines"),
        ("Developer / Service Provider", "______________________________"),
        ("Project", "Court Management System for club operations, player queues, live display, finance ledger, and reservations"),
        ("Contract Price", "PHP 35,000 one-time system fee; PHP 2,500 yearly domain; PHP 2,000 monthly maintenance"),
    ]
    for i, (label, value) in enumerate(rows):
        c0, c1 = meta.rows[i].cells
        set_cell_shading(c0, "F2F4F7")
        for c in (c0, c1):
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            c.paragraphs[0].paragraph_format.space_after = Pt(2)
        style_run(c0.paragraphs[0].add_run(label), bold=True, color="1F4D78")
        c1.paragraphs[0].add_run(value)

    doc.add_paragraph()
    doc.add_paragraph(
        "This System Development and Implementation Agreement (the \"Agreement\") is entered into by and between the Client and the Developer / Service Provider identified above. The purpose of this Agreement is to set out the scope, fees, responsibilities, deliverables, and acceptance terms for the HAFF Leisure Club Court Management System."
    )

    add_clause(doc, "1. Scope of Work", [
        "The Developer shall provide, configure, and assist with the implementation of a web-based court management system for HAFF Leisure Club. The system is intended to support day-to-day club operations and open-play court flow.",
        [
            "Administrative screen for staff session control, player check-in, queue management, court assignments, and match completion.",
            "Player-facing screen for members to view queue status, wait time, and related club activity.",
            "Venue TV display for live court board and real-time court rotation visibility.",
            "Finance screen or ledger for club revenue tracking and administrative reference.",
            "Court reservation calendar for scheduling and operational planning.",
            "Supabase-backed database, realtime updates, and storage configuration as required for the deployed system.",
            "Deployment support for Vercel hosting and Supabase project configuration.",
        ],
    ])

    add_clause(doc, "2. Deliverables", [
        "The Developer shall deliver a functioning production-ready web application based on the existing HAFF Leisure Club codebase, together with the deployment configuration and reasonable handover support needed for the Client to operate the system.",
        [
            "Live web application deployed to the agreed hosting account or environment.",
            "Source code repository or project files for the delivered system.",
            "Basic deployment notes covering Vercel, Supabase, and required environment variables.",
            "Initial setup assistance for administrator access, courts, and operating data.",
        ],
    ])

    add_clause(doc, "3. Fees and Payment Terms", [
        "The commercial terms for this Agreement are based on the Client's accepted counteroffer: PHP 35,000 as a one-time system fee, PHP 2,500 per year for the domain, and PHP 2,000 per month for maintenance support, unless the parties agree in writing to different terms.",
    ])
    pay = doc.add_table(rows=5, cols=3)
    pay.style = "Table Grid"
    set_table_width(pay, [2880, 2520, 3720])
    set_cell_margins(pay)
    set_borders(pay)
    for j, head in enumerate(("Item", "Amount", "Due")):
        cell = pay.rows[0].cells[j]
        set_cell_shading(cell, "F2F4F7")
        style_run(cell.paragraphs[0].add_run(head), bold=True, color="1F4D78")
    data = [
        ("System development and implementation", "PHP 35,000", "One-time payment; due upon signing, launch, or another written date agreed by both parties"),
        ("Domain", "PHP 2,500 / year", "Yearly domain cost, subject to registrar renewal rules and the final domain selected"),
        ("Maintenance support", "PHP 2,000 / month", "Monthly; covers reasonable technical maintenance support and routine system assistance"),
        ("Future enhancements, major redesigns, new modules, hardware, SMS, paid email, or other third-party services", "Not included", "Subject to separate written quotation or paid directly by the Client"),
    ]
    for i, row in enumerate(data, start=1):
        for j, text in enumerate(row):
            pay.rows[i].cells[j].paragraphs[0].add_run(text)

    add_clause(doc, "4. Commercial Comparison", [
        "For reference, the Client reviewed an alternative proposal with PHP 35,000 setup, PHP 8,500 custom domain/hosting, and PHP 4,000 monthly retainership pricing. That alternative totals approximately PHP 91,500 for the first year using the custom-domain option.",
        "Under this Agreement, the first-year commercial total is approximately PHP 61,500, consisting of PHP 35,000 system fee, PHP 2,500 yearly domain cost, and twelve months of PHP 2,000 maintenance. This reflects approximately PHP 30,000 in first-year savings compared with the alternative custom-domain package, and PHP 24,000 per year in lower recurring maintenance cost.",
        "If comparing against a PHP 90,000 one-time or bundled system package, the PHP 35,000 system fee is approximately PHP 55,000 lower before domain and maintenance costs."
    ])

    add_clause(doc, "5. Client Responsibilities", [
        "The Client shall provide timely access, information, approvals, and decisions necessary for completion and deployment of the system.",
        [
            "Provide accurate club, court, player, pricing, and operations information.",
            "Provide or approve access to hosting, Supabase, domain, and related accounts where needed.",
            "Review the system promptly and report issues with clear reproduction steps.",
            "Ensure authorized staff use the system responsibly and protect account credentials.",
        ],
    ])

    add_clause(doc, "6. Timeline and Acceptance", [
        "The parties shall agree on the target launch date separately. The system shall be considered accepted when it is deployed and the core workflows described in Section 1 can be performed in the agreed environment, except for minor issues that do not materially prevent use.",
        "Any requested changes outside the agreed scope shall be treated as change requests and may affect timing and fees."
    ])

    add_clause(doc, "7. Change Requests", [
        "A change request includes any new feature, major workflow revision, redesign, additional integration, data migration beyond the agreed setup, or repeated revision caused by changed Client instructions. Change requests must be approved in writing before work begins."
    ])

    add_clause(doc, "8. Maintenance and Support", [
        "The PHP 35,000 one-time system fee covers development and implementation of the agreed system. The PHP 2,000 monthly maintenance fee covers reasonable technical maintenance support, routine updates, and basic assistance after launch. Emergency support, major new feature work, hardware installation, server migrations, and third-party platform issues are not included unless separately agreed in writing."
    ])

    add_clause(doc, "9. Ownership and License", [
        "Upon full payment, the Client receives the right to use the delivered HAFF Leisure Club system for its internal club operations. Unless otherwise agreed in writing, reusable developer tools, general methods, libraries, prior work, and non-client-specific know-how remain the property of the Developer.",
        "The system may include open-source or third-party components subject to their own licenses and service terms."
    ])

    add_clause(doc, "10. Confidentiality and Data", [
        "Each party shall keep confidential information received from the other party confidential and shall use it only for purposes of this Agreement. The Client remains responsible for the accuracy, legality, retention, and appropriate use of player, member, financial, and operational data entered into the system.",
        "The Developer shall not intentionally disclose Client data except as required to perform the work, comply with law, or support accounts and services approved by the Client."
    ])

    add_clause(doc, "11. Warranty and Limitations", [
        "The Developer shall use reasonable skill and care in delivering the system. The Developer does not guarantee uninterrupted operation of third-party services such as hosting, database, network, domain, browser, device, payment, SMS, email, or other external platforms.",
        "To the fullest extent allowed by law, the Developer shall not be liable for indirect, incidental, special, consequential, or lost-profit damages. The Developer's total liability under this Agreement shall not exceed the amount actually paid by the Client under this Agreement."
    ])

    add_clause(doc, "12. General Terms", [
        "This Agreement may be amended only by written agreement of both parties. If any provision is found unenforceable, the remaining provisions shall continue in effect. This Agreement represents the parties' full understanding regarding the system described above."
    ])

    add_signature_block(doc)
    doc.save(OUT)


if __name__ == "__main__":
    build()
