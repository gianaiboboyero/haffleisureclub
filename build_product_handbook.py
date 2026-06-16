from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.section import WD_SECTION

OUT = "/Users/mbam22022/Documents/Haff Leisure Club/HAFF_Leisure_Club_Product_Handbook.docx"

GREEN = "12372C"
DEEP = "082D23"
BRASS = "CBEF43"
GOLD = "F4C95D"
IVORY = "F7F3E8"
MUTED = "5D7169"
LIGHT = "EDF2ED"
WHITE = "FFFFFF"

doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.8)
section.bottom_margin = Inches(0.8)
section.left_margin = Inches(0.85)
section.right_margin = Inches(0.85)
section.header_distance = Inches(0.4)
section.footer_distance = Inches(0.4)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Aptos"
normal.font.size = Pt(10.5)
normal.font.color.rgb = RGBColor.from_string(GREEN)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.15

for name, size, color, before, after in [
    ("Title", 29, GREEN, 0, 8),
    ("Subtitle", 13, MUTED, 0, 12),
    ("Heading 1", 18, GREEN, 18, 8),
    ("Heading 2", 14, GREEN, 14, 6),
    ("Heading 3", 11.5, MUTED, 10, 4),
]:
    st = styles[name]
    st.font.name = "Aptos Display" if name != "Normal" else "Aptos"
    st.font.size = Pt(size)
    st.font.color.rgb = RGBColor.from_string(color)
    st.font.bold = name != "Subtitle"
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), fill)

def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")

def add_run(p, text, bold=False, color=GREEN, size=None, italic=False):
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.name = "Aptos"
    r.font.color.rgb = RGBColor.from_string(color)
    if size:
        r.font.size = Pt(size)
    return r

def para(text="", bold_prefix=None, style=None):
    p = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        add_run(p, bold_prefix, bold=True)
        add_run(p, text[len(bold_prefix):])
    else:
        add_run(p, text)
    return p

def bullets(items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.18)
        p.paragraph_format.space_after = Pt(4)
        add_run(p, item)

def numbered(items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.first_line_indent = Inches(-0.2)
        p.paragraph_format.space_after = Pt(5)
        add_run(p, item)

def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    for i, header in enumerate(headers):
        c = t.rows[0].cells[i]
        shade(c, GREEN)
        set_cell_margins(c)
        c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = c.paragraphs[0]
        add_run(p, header, bold=True, color=WHITE, size=9.5)
        if widths: c.width = Inches(widths[i])
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            c = cells[i]
            set_cell_margins(c)
            if len(t.rows) % 2 == 1: shade(c, LIGHT)
            p = c.paragraphs[0]
            add_run(p, str(val), size=9)
            if widths: c.width = Inches(widths[i])
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return t

def callout(title, text, fill=LIGHT):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = t.cell(0, 0)
    shade(c, fill)
    set_cell_margins(c, 160, 180, 160, 180)
    p = c.paragraphs[0]
    add_run(p, title + "\n", bold=True, color=GREEN, size=11)
    add_run(p, text, color=GREEN, size=10)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)

def page_break():
    doc.add_page_break()

header = section.header
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
add_run(hp, "HAFF LEISURE CLUB | PRODUCT HANDBOOK", bold=True, color=MUTED, size=8)
footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(fp, "Cadiz City | haffleisureclub.com | June 2026", color=MUTED, size=8)

# Cover
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(74)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, "HAFF LEISURE CLUB", bold=True, color=GREEN, size=12)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(8)
add_run(p, "Product & Experience Handbook", bold=True, color=GREEN, size=30)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, "A complete guide to the club website, operational system, member community, and future product direction", color=MUTED, size=13)
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(32)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, "CADIZ CITY", bold=True, color=GOLD, size=11)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, "Current-state documentation based on the live implementation as of June 14, 2026", color=MUTED, size=9.5, italic=True)
doc.add_paragraph()
callout("The product in one sentence", "HAFF Leisure Club is a mobile-friendly club operating system that connects public discovery, member accounts, open-play rotation, court operations, community conversation, feedback, play history, finance, scheduling, and venue TV displays in one branded experience.")

page_break()
doc.add_heading("1. Executive Overview", level=1)
para("HAFF Leisure Club is not merely a marketing website. It is an evolving digital operating layer for a physical leisure venue in Cadiz City. It serves three connected purposes: attracting and informing visitors, helping members participate in club life, and helping staff operate open play, courts, schedules, player records, and financial activity.")
para("The experience is designed around the real rhythm of a recreation club. A visitor learns what HAFF is and joins. A member signs in once and gains access to community and player tools. Staff manage the flow of people and courts. A shared TV keeps everyone oriented inside the venue.")
callout("Core product promise", "Make club participation feel simple for members and make high-volume open-play operations manageable for staff, even when connectivity is imperfect.")

doc.add_heading("Who the product serves", level=2)
table(["Audience", "Primary need", "Product response"], [
    ("Visitor / prospective member", "Understand HAFF and decide whether to join", "Home experience, live venue context, testimonials, registration"),
    ("Member / player", "Know when to play, communicate, manage identity, and review activity", "Prisma account, player dashboard, community, recap"),
    ("Front desk / coordinator", "Keep players, queues, courts, and sessions moving", "Admin rotation controls, player and court management"),
    ("Club owner / manager", "Oversee operations, finance, schedules, and member sentiment", "Admin, finance, calendar, testimonials, anonymous reports"),
    ("In-venue audience", "See what is happening without asking staff", "Responsive TV display, court status, timers, queue groups, announcements"),
], [1.35, 2.15, 2.8])

doc.add_heading("Current product principles", level=2)
bullets([
    "One HAFF account: email/password authentication links a User to a Player profile.",
    "Mobile first: the most common member interactions are designed for phones.",
    "Operational clarity: court state, queue state, and timing should be visible and actionable.",
    "Offline tolerance: operational data is stored locally first and synchronized when possible.",
    "Community with moderation: conversation and testimonials are member-driven but administratively reviewable.",
    "Brand consistency: dark court green, warm ivory, and bright brass/lime accents establish a recognizable Cadiz City identity.",
])

page_break()
doc.add_heading("2. Information Architecture & Route Map", level=1)
para("The product uses client-side routes rendered by a React/Vite single-page application. The canonical home route is /home. Older /landing URLs are normalized to /home for backward compatibility.")
table(["Route", "Audience", "Purpose", "Access"], [
    ("/home", "Everyone", "Public discovery, sign in/register entry, live context, feedback, testimonials", "Public"),
    ("/community", "Members", "Authentication, general chat, feedback, testimonials, play recap", "Public gate; tools require login"),
    ("/player", "Members", "Personal queue, check-in, park mode, profile, play status", "Prisma login required"),
    ("/admin", "Staff", "Open-play control center and operational administration", "Admin role required"),
    ("/calendar", "Staff", "Court reservations and schedule management", "Admin role required"),
    ("/finance", "Staff", "Transactions, revenue, and payment management", "Admin role required"),
    ("/tv", "Venue", "Large-screen court, timer, queue, and announcement display", "Operational display"),
], [1.0, 1.0, 3.15, 1.35])
para("Navigation changes by context. Regular players see a reduced set of choices to avoid exposing staff tools. Authenticated administrators receive the full dock: Home, Admin, Players, Community, Calendar, Finance, and TV Display. The Home footer contains a discreet Staff access entry that opens the admin login.")

doc.add_heading("Global account behavior", level=2)
numbered([
    "At application startup or refresh, the browser checks the Prisma-backed session before revealing the interface.",
    "Signed-out visitors see Sign in and Register in the Home header.",
    "Signed-in users see a stable upper-right account control with avatar and display name.",
    "The account menu contains Sign out. Signing out clears the server session cookie and local player link, then returns the user to Home.",
    "The same session applies across Community, Player, Admin, Calendar, and Finance; users should not authenticate separately for each surface.",
])

page_break()
doc.add_heading("3. Public Home Experience", level=1)
para("Home is the front door of the club. Its role is to explain HAFF quickly, establish the premium sport-and-leisure atmosphere, and direct people toward the action that makes sense for them.")
doc.add_heading("What visitors encounter", level=2)
bullets([
    "Cadiz City brand header with HAFF identity and familiar Sign in/Register actions.",
    "Hero messaging positioning HAFF as a sanctuary of sport and leisure.",
    "Live or locally synchronized club statistics, including checked-in players and court activity.",
    "Clear entry cards for player participation and the venue TV display.",
    "An explanation of the open-play rotation: check in, form stacks, and play the next available court.",
    "Venue lifestyle content describing courts, lounge, cafe, and social atmosphere.",
    "Anonymous improvement reporting without requiring an account.",
    "Approved member testimonials, controlled through admin moderation.",
    "A discreet Staff access link in the footer instead of a prominent public admin button.",
])
doc.add_heading("Ideal visitor journey", level=2)
numbered([
    "A visitor arrives at /home from social media, search, a QR code, or a direct link.",
    "They understand the venue and open-play concept within the first screen.",
    "They choose Register to create a member account or Sign in if they already belong.",
    "After authentication, they can enter Community and Player without another login.",
    "They can later return to Home and remain recognized through the 30-day session.",
])
callout("Experience goal", "The public site should feel like entering the physical club: energetic but orderly, social but premium, and immediately clear about what to do next.", "FFF8E8")

page_break()
doc.add_heading("4. Member Registration, Login & Identity", level=1)
para("Authentication is implemented directly with Prisma and PostgreSQL. Registration creates both a User account and a linked Player profile. Passwords are not stored directly; they are salted and hashed using scrypt.")
table(["Stage", "User action", "System behavior"], [
    ("Register", "Enter display name, skill level, email, and password", "Creates User + Player records and starts a session"),
    ("Sign in", "Enter email and password", "Verifies hash and issues an HttpOnly session cookie"),
    ("Refresh", "Reload any route", "Validates session before rendering the app"),
    ("Open Player", "Visit /player", "Loads the Player linked to the signed-in User"),
    ("Sign out", "Use upper-right account menu", "Deletes session and clears local player identity"),
], [1.0, 2.15, 3.35])
doc.add_heading("Identity model", level=2)
bullets([
    "User stores credentials, role, status, and session relationships.",
    "Player stores the club-facing identity: display name, skill level, rating, avatar, contact data, tags, and play statistics.",
    "User.playerId is unique, establishing one linked player identity per account.",
    "Roles currently use ADMIN and MEMBER values.",
    "Account status can prevent inactive users from resolving a valid session.",
])
doc.add_heading("UX considerations", level=2)
bullets([
    "Home is the primary authentication entry; Player should never present a competing phone-code login.",
    "The Player route explains that login is required and links directly to Sign in/Register.",
    "Saved browser credentials are styled to remain white/ivory rather than showing intrusive gray autofill fields.",
    "The account menu should remain stable after refresh and across routes.",
])

page_break()
doc.add_heading("5. Player Experience", level=1)
para("The Player dashboard answers one urgent question: what should I do now? It translates operational club state into a personal view.")
doc.add_heading("Player capabilities", level=2)
bullets([
    "See whether the player is checked in, parked, waiting, reserved, or currently playing.",
    "Estimate queue position and expected wait based on current stacks, courts, and timers.",
    "Check in and manage park/resume behavior.",
    "View a compact TV-style preview of current club activity.",
    "Edit profile information, play style, skill level, emergency note, and avatar.",
    "Display a member QR pass.",
    "Review player metrics such as total games and days played.",
])
doc.add_heading("Possible player scenarios", level=2)
table(["Scenario", "Expected experience"], [
    ("Arriving at the club", "Sign in once, open Player, tap check-in, receive immediate status"),
    ("Taking a break", "Use park mode without losing account identity"),
    ("Waiting for a court", "See position, grouped stack, and approximate timing"),
    ("Called to play", "See reserved court and roster reflected in personal and TV views"),
    ("After play", "Match completion updates history and contributes to recap metrics"),
    ("Returning another day", "Session remains recognized until sign-out or expiry"),
], [1.8, 4.7])
callout("Current limitation", "The operational player state is partly offline-first browser data, while account/community data is server-backed. A future phase should make Prisma the authoritative source for all player operations across devices.")

page_break()
doc.add_heading("6. Community & Social Experience", level=1)
para("Community turns HAFF from a utility into a club network. It combines conversation, feedback, testimonials, and activity storytelling under one member identity.")
doc.add_heading("General chat", level=2)
bullets([
    "Club-wide messenger-style room for authenticated members.",
    "Messages show member avatar, display name, role indicator, time, date grouping, and edit state.",
    "Replies preserve the context of the referenced message.",
    "Members can react with emoji, edit their own messages, remove their own messages, and report others.",
    "Admins can remove messages and review reported content.",
    "Messages appear optimistically on send and refresh from the server every second while the page is visible.",
    "Members can sort the conversation oldest-first or newest-first.",
])
doc.add_heading("Interaction direction", level=2)
para("The desired next evolution is a Facebook Messenger-like reaction experience: press-and-hold on mobile or hover/click on desktop to reveal a compact reaction tray, with a lightweight animation and clear selected state. This should replace permanently visible reaction/action controls that make each bubble feel crowded.")
doc.add_heading("Testimonials", level=2)
bullets([
    "Authenticated members submit a quote and rating.",
    "Submissions begin as PENDING.",
    "Admins approve or reject them.",
    "Approved testimonials may appear publicly on Home.",
])
doc.add_heading("Anonymous improvement reports", level=2)
bullets([
    "Visitors and members can submit without attaching identity.",
    "Category options include facilities, courts, scheduling, app, staff/service, safety, and other.",
    "Contact information is optional.",
    "A hashed source/day value supports basic rate limiting without storing a raw IP address.",
    "Admins can review and mark reports as read.",
])

page_break()
doc.add_heading("7. Play Recap & Shareable Activity", level=1)
para("The recap feature borrows the social storytelling pattern popularized by fitness platforms: turn raw participation data into a compact achievement card members can copy, download, or share.")
doc.add_heading("Current recap metrics", level=2)
bullets([
    "Active play duration derived from completed matches.",
    "Total completed games.",
    "Wins based on team membership and final score.",
    "Unique players encountered.",
    "Number of courts played.",
    "Lifetime player totals where available.",
])
doc.add_heading("Sharing options", level=2)
bullets([
    "Copy a formatted text summary to the clipboard.",
    "Use the device-native share sheet when supported.",
    "Generate and download a portrait PNG recap card.",
    "Include HAFF Leisure Club - Cadiz City branding and website attribution.",
])
callout("Product opportunity", "Recaps can become a retention loop: weekly summaries, personal bests, streaks, most-played partners, skill progression, and branded templates for Stories or group chats.")

page_break()
doc.add_heading("8. Admin Operations", level=1)
para("The Admin dashboard is the operational center of the product. It is protected by the same Prisma session system and requires the ADMIN role.")
doc.add_heading("Open-play control", level=2)
bullets([
    "View checked-in players and active lounge participants.",
    "Build and reorder player stacks.",
    "Generate balanced teams based on player skill and rating.",
    "Reserve, start, clear, and finish courts.",
    "Move players between queue and courts.",
    "Manage overtime thresholds and match duration.",
    "Publish club status messages and trigger sound/voice announcements.",
])
doc.add_heading("Player administration", level=2)
bullets([
    "Create, edit, activate, deactivate, and remove player records.",
    "Manage identity, phone, email, access metadata, skill level, rating, tags, and photos.",
    "Review player totals and participation history.",
])
doc.add_heading("Court and session administration", level=2)
bullets([
    "Create and update courts, numbers, notes, and availability.",
    "Start and end open-play sessions.",
    "Associate courts and checked-in players with the current session.",
    "Review completed match history and scores.",
])
doc.add_heading("Community moderation", level=2)
bullets([
    "Approve or reject testimonials.",
    "Review anonymous improvement reports.",
    "Review reported chat messages and remove inappropriate content.",
])

page_break()
doc.add_heading("9. Calendar, Finance & TV", level=1)
doc.add_heading("Calendar and reservations", level=2)
para("The Calendar surface is designed for court availability and booking. Staff can choose a court, player, start and end time, and manage reservation status. This area is currently an administrative tool rather than a public self-service booking flow.")
doc.add_heading("Finance", level=2)
para("The Finance surface tracks transactions and revenue-related activity. It supports transaction creation, completion, and payment-oriented club workflows. It is role-protected and appears in the full administrator navigation.")
doc.add_heading("Venue TV display", level=2)
bullets([
    "Responsive landscape-first display for large venue screens.",
    "Shows courts, current matches, teams, scores, and countdown timers.",
    "Displays waiting queue groups and reserved players.",
    "Highlights overtime courts.",
    "Supports club-wide status announcements.",
    "Uses audio announcements for overtime and next-player calls when sound is enabled.",
    "Includes layouts for desktop TVs, tablets, mobile/portrait displays, and foldable-sized viewports.",
])
callout("Operational value", "The TV reduces repeated questions at the front desk. It makes the rotation system visible, creates trust in queue fairness, and gives the venue a more organized atmosphere.", "FFF8E8")

page_break()
doc.add_heading("10. Data & Technical Architecture", level=1)
table(["Layer", "Technology", "Responsibility"], [
    ("Web application", "React 18 + TypeScript + Vite", "Routes, UI, client behavior, PWA"),
    ("Styling", "Tailwind CSS 3 + custom CSS", "Responsive layout, HAFF visual system"),
    ("Server API", "Vercel serverless functions", "Auth, chat, feedback, testimonials, recap, sync"),
    ("Database", "Prisma ORM + PostgreSQL", "Accounts, sessions, players, operations, community data"),
    ("Offline storage", "Dexie / IndexedDB", "Local-first club operational state and sync queue"),
    ("Deployment", "Vercel", "Production hosting, functions, domain, HTTPS"),
    ("Domain/DNS", "Namecheap + Vercel DNS targets", "haffleisureclub.com and www redirect"),
], [1.25, 2.05, 3.2])
doc.add_heading("Prisma data domains", level=2)
table(["Domain", "Models"], [
    ("Identity", "User, AuthSession, Player"),
    ("Operations", "Court, Session, Match"),
    ("Offline synchronization", "SyncEvent"),
    ("Communication", "Notification, ChatMessage, ChatReaction, ChatReport"),
    ("Voice of member", "Testimonial, ImprovementReport"),
], [2.0, 4.5])
doc.add_heading("Offline-first behavior", level=2)
para("Operational actions are saved into IndexedDB first, with a synchronization queue used to push changes to the server. This is useful for a venue where Wi-Fi may be inconsistent. The interface exposes online/offline and pending-sync state. Courts are seeded when no local courts exist, but deleted/demo players are not automatically recreated.")
doc.add_heading("Current real-time model", level=2)
para("The production serverless environment does not maintain a persistent Socket.IO connection. Most operational freshness uses polling or explicit synchronization. General chat polls every second while visible and uses optimistic rendering for immediate send feedback.")

page_break()
doc.add_heading("11. Security, Privacy & Moderation", level=1)
bullets([
    "Passwords use scrypt with a random salt before storage.",
    "Session tokens are random, stored only as SHA-256 hashes in the database, and delivered through Secure, HttpOnly, SameSite=Lax cookies.",
    "Sessions currently last up to 30 days.",
    "Admin APIs verify the server-side role rather than trusting a browser flag.",
    "Chat editing/deletion checks ownership or administrator role.",
    "Anonymous reports avoid storing raw source identity and enforce basic per-source rate limiting.",
    "Messages are soft-deleted, preserving moderation/history structure while hiding content.",
])
doc.add_heading("Risks and improvements", level=2)
table(["Area", "Current concern", "Recommended improvement"], [
    ("Registration", "Open registration can attract spam", "Email verification, CAPTCHA/Turnstile, invite or QR enrollment options"),
    ("Password recovery", "No self-service reset flow", "Secure reset tokens and email delivery"),
    ("Authorization", "Some legacy operational APIs are broadly accessible", "Enforce role checks on every mutation endpoint"),
    ("Audit trail", "Limited admin action history", "Add immutable moderation and operations audit events"),
    ("Chat abuse", "Basic reporting only", "Block/mute, escalation states, rate controls, moderation notes"),
    ("Data consistency", "Local operational state and server data can diverge", "Server-authoritative event model with conflict handling"),
], [1.15, 2.55, 2.8])

page_break()
doc.add_heading("12. End-to-End User Journeys", level=1)
doc.add_heading("Journey A: New member", level=2)
numbered([
    "Discover HAFF through /home and understand the club experience.",
    "Choose Register in the upper-right Home header.",
    "Create an account with display name, skill level, email, and password.",
    "Land in Community as a recognized member.",
    "Open Player and see the linked personal dashboard without another login.",
    "Check in at the venue, follow queue status, and play.",
    "Later submit a testimonial or share a play recap.",
])
doc.add_heading("Journey B: Returning player", level=2)
numbered([
    "Open /home or scan a club QR.",
    "Existing session is resolved before the page appears.",
    "Use the upper-right avatar menu for account context.",
    "Open Player, check status, and check in.",
    "Join Community chat while waiting.",
    "Receive court information through Player and TV.",
])
doc.add_heading("Journey C: Front-desk administrator", level=2)
numbered([
    "Use Staff access in the Home footer and authenticate with the admin account.",
    "Open Admin and start or resume the current session.",
    "Check in players, build stacks, and monitor lounge volume.",
    "Assign a stack to an available court and start the timer.",
    "Use TV and voice announcements to call players.",
    "Finish the match, record scores, and rotate the next group.",
    "Review feedback, community reports, reservations, and finance as needed.",
])

page_break()
doc.add_heading("13. Experience Quality & Accessibility", level=1)
bullets([
    "Touch targets generally use a minimum height near 44 pixels.",
    "Mobile navigation is horizontally scrollable when administrator options exceed the viewport.",
    "Forms use labels/placeholders, explicit autofill behavior, and high-contrast input treatment.",
    "Reduced-motion preferences disable or minimize animations.",
    "The global loader uses the same textured/grid HAFF canvas as Home and dashboard pages.",
    "The loader text is intentionally branded as “loadink” and displayed without letter spacing.",
    "Responsive TV layouts adapt to narrow and short screens.",
    "Account state is resolved before rendering to reduce visual popping after refresh.",
])
doc.add_heading("UX areas to continue refining", level=2)
bullets([
    "Add proper focus trapping and outside-click behavior to the account menu.",
    "Replace browser prompt/confirm interactions in chat with designed dialogs.",
    "Add accessible labels to every icon-only control and announce chat send failures.",
    "Use skeletons inside content regions where full-screen loading is unnecessary.",
    "Test keyboard-only navigation, screen readers, Android autofill, iOS safe areas, and foldable devices.",
    "Standardize date/time formatting for Philippine locale and club timezone.",
])

page_break()
doc.add_heading("14. Product Roadmap Recommendations", level=1)
table(["Phase", "Priority", "Recommended outcomes"], [
    ("1. Reliability", "Immediate", "Unify Prisma and operational state, secure all mutations, password reset, email verification, error monitoring"),
    ("2. Member UX", "Near term", "Messenger reaction tray, notification center, profile completion, stronger recap cards, attendance history"),
    ("3. Operations", "Near term", "Server-authoritative queue, live court events, staff audit trail, reservation conflict prevention"),
    ("4. Growth", "Medium term", "Membership plans, self-service booking, payments, referral links, event registration"),
    ("5. Intelligence", "Long term", "Demand forecasts, skill recommendations, fair-play analytics, retention insights, automated weekly recaps"),
], [1.15, 1.1, 4.25])
doc.add_heading("Highest-value next improvements", level=2)
numbered([
    "Make Prisma/PostgreSQL authoritative for live player, queue, court, match, reservation, and finance data.",
    "Replace one-second chat polling with managed real-time delivery such as database change subscriptions or a hosted realtime service compatible with Prisma.",
    "Build Facebook-style message interactions: hold/hover reaction tray, contextual action sheet, read position, and unread count.",
    "Add email verification, password reset, and account recovery.",
    "Introduce proper membership and payment rules before opening public booking.",
    "Create analytics for court utilization, wait time, repeat participation, and member satisfaction.",
])

page_break()
doc.add_heading("15. Product Definition Summary", level=1)
para("HAFF Leisure Club is best understood as a connected venue platform with four faces:")
table(["Face", "Definition"], [
    ("Brand website", "Explains the club, its atmosphere, and how to participate."),
    ("Member app", "Connects identity, player status, community, feedback, and play history."),
    ("Operations console", "Coordinates players, stacks, sessions, courts, reservations, and money."),
    ("Venue display system", "Broadcasts the live state of play to everyone in the building."),
], [1.55, 4.95])
para("Its strongest differentiator is that these surfaces are not separate products. They share one brand, one account system, and one operational story: arrive, join, wait fairly, play, connect, and return.")
callout("North-star experience", "A member should be able to arrive at HAFF, understand what is happening, join play, communicate, and leave with a record of the experience—all without needing to ask staff what to do next.")

doc.add_heading("Appendix A: Current API Surface", level=1)
table(["Endpoint area", "Key actions"], [
    ("Authentication", "me, register, login, logout"),
    ("Community", "messages, send, edit, delete, react, report, moderation reports"),
    ("Testimonials", "approved, submit, pending, moderate"),
    ("Feedback", "submit, list, read"),
    ("Recap", "authenticated player recap"),
    ("Operations", "players GET, courts GET, sync POST"),
], [1.7, 4.8])

doc.add_heading("Appendix B: Terminology", level=1)
table(["Term", "Meaning"], [
    ("Stack", "A waiting group of players intended for the next available match."),
    ("Park mode", "A temporary pause that keeps the player identity available without placing them in the active queue."),
    ("Reserved court", "A court assigned to an upcoming player group before play starts."),
    ("Open play", "Club format where players rotate through available courts rather than holding a private booking."),
    ("Play recap", "A shareable summary of completed match activity."),
    ("Local-first", "Saving operational changes in the browser before synchronizing to the server."),
], [1.45, 5.05])

doc.core_properties.title = "HAFF Leisure Club Product & Experience Handbook"
doc.core_properties.subject = "Website purpose, flows, features, architecture, user experience, and roadmap"
doc.core_properties.author = "HAFF Leisure Club"
doc.save(OUT)
print(OUT)
